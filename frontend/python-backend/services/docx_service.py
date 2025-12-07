from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path
import tempfile

class DocxService:
    """Service for creating and manipulating DOCX documents"""

    @staticmethod
    def create_document(content: dict) -> Path:
        """
        Create a new DOCX document from content dictionary.

        Args:
            content: Dictionary with structure:
                {
                    "title": "Document Title",
                    "sections": [
                        {"heading": "Section 1", "body": "Content..."},
                        {"heading": "Section 2", "body": "Content..."}
                    ]
                }

        Returns:
            Path to created DOCX file
        """
        doc = Document()

        # Add title if provided
        if content.get("title"):
            title = doc.add_heading(content["title"], level=0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add sections
        for section in content.get("sections", []):
            if section.get("heading"):
                doc.add_heading(section["heading"], level=1)

            if section.get("body"):
                # Split body into paragraphs
                paragraphs = section["body"].split("\n\n")
                for para_text in paragraphs:
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())

        # Save to temp file
        temp_dir = Path(tempfile.gettempdir()) / "docx-creator"
        temp_dir.mkdir(exist_ok=True)
        output_path = temp_dir / f"document_{tempfile._get_candidate_names().__next__()}.docx"

        doc.save(str(output_path))
        return output_path

    @staticmethod
    def create_from_template(template_type: str, content: dict) -> Path:
        """
        Create a document from a predefined template.

        Args:
            template_type: Type of template (report, memo, letter, etc.)
            content: Content to fill in template

        Returns:
            Path to created DOCX file
        """
        doc = Document()

        if template_type == "report":
            return DocxService._create_report(doc, content)
        elif template_type == "memo":
            return DocxService._create_memo(doc, content)
        elif template_type == "letter":
            return DocxService._create_letter(doc, content)
        else:
            # Default to basic document
            return DocxService.create_document(content)

    @staticmethod
    def _create_report(doc: Document, content: dict) -> Path:
        """Create a professional report"""
        # Title page
        title = doc.add_heading(content.get("title", "Report"), level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add date and author
        doc.add_paragraph()
        if content.get("date"):
            date_para = doc.add_paragraph(f"Date: {content['date']}")
            date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if content.get("author"):
            author_para = doc.add_paragraph(f"Prepared by: {content['author']}")
            author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_page_break()

        # Executive Summary
        if content.get("executive_summary"):
            doc.add_heading("Executive Summary", level=1)
            doc.add_paragraph(content["executive_summary"])
            doc.add_page_break()

        # Main sections
        for section in content.get("sections", []):
            if section.get("heading"):
                doc.add_heading(section["heading"], level=1)

            if section.get("body"):
                paragraphs = section["body"].split("\n\n")
                for para_text in paragraphs:
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())

        # Save
        temp_dir = Path(tempfile.gettempdir()) / "docx-creator"
        temp_dir.mkdir(exist_ok=True)
        output_path = temp_dir / f"report_{tempfile._get_candidate_names().__next__()}.docx"
        doc.save(str(output_path))
        return output_path

    @staticmethod
    def _create_memo(doc: Document, content: dict) -> Path:
        """Create a business memo"""
        # Memo header
        doc.add_heading("MEMORANDUM", level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()

        # Memo fields
        doc.add_paragraph(f"TO: {content.get('to', '[Recipient]')}")
        doc.add_paragraph(f"FROM: {content.get('from', '[Sender]')}")
        doc.add_paragraph(f"DATE: {content.get('date', '[Date]')}")
        doc.add_paragraph(f"RE: {content.get('subject', '[Subject]')}")
        doc.add_paragraph()
        doc.add_paragraph("_" * 70)
        doc.add_paragraph()

        # Body
        if content.get("body"):
            paragraphs = content["body"].split("\n\n")
            for para_text in paragraphs:
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())

        # Save
        temp_dir = Path(tempfile.gettempdir()) / "docx-creator"
        temp_dir.mkdir(exist_ok=True)
        output_path = temp_dir / f"memo_{tempfile._get_candidate_names().__next__()}.docx"
        doc.save(str(output_path))
        return output_path

    @staticmethod
    def _create_letter(doc: Document, content: dict) -> Path:
        """Create a business letter"""
        # Sender address
        if content.get("sender_address"):
            doc.add_paragraph(content["sender_address"])
        doc.add_paragraph()

        # Date
        if content.get("date"):
            doc.add_paragraph(content["date"])
        doc.add_paragraph()

        # Recipient address
        if content.get("recipient_address"):
            doc.add_paragraph(content["recipient_address"])
        doc.add_paragraph()

        # Salutation
        salutation = content.get("salutation", "Dear Sir/Madam,")
        doc.add_paragraph(salutation)
        doc.add_paragraph()

        # Body
        if content.get("body"):
            paragraphs = content["body"].split("\n\n")
            for para_text in paragraphs:
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())
        doc.add_paragraph()

        # Closing
        closing = content.get("closing", "Sincerely,")
        doc.add_paragraph(closing)
        doc.add_paragraph()
        doc.add_paragraph()

        # Signature
        if content.get("signature"):
            doc.add_paragraph(content["signature"])

        # Save
        temp_dir = Path(tempfile.gettempdir()) / "docx-creator"
        temp_dir.mkdir(exist_ok=True)
        output_path = temp_dir / f"letter_{tempfile._get_candidate_names().__next__()}.docx"
        doc.save(str(output_path))
        return output_path
